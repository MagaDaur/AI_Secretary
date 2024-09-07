import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import win32com.client
import os

doc=Document()

paragraph = doc.add_paragraph()
password="12345"
isPassword=True

run = paragraph.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
theme="такой-то"
date=datetime.datetime.now()
cur=str(date.day)+"."+str(date.month)+"."+str(date.year)
months = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
           'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
count_of_members=3
count_of_questions=2
voc={1:"I",2:"II",3:"III",4:"IV",5:"V",
     6:"VI",7:"VII",8:"VIII",9:"IX",10:"X",
     11:"XI",12:"XII",13:"XIII",14:"XIV",15:"XV",
     16:"XVI",17:"XVII",18:"XVIII",19:"XIX",20:"XX"}
questions={
    "Название вопроса 1":{
        "speakers":["Докладчик 1", "Докладчик 2"],
        "decissions":{
            "Принятое решение 1":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            },
            "Принятое решение 2":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            }
        }
    },
    "Название вопроса 2":{
        "speakers":["Докладчик 1", "Докладчик 2", "Докладчик 3"],
        "decissions":{
            "Принятое решение 1":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            },
            "Принятое решение 2":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            },
            "Принятое решение 3":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            },
            "Принятое решение 4":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            }
        }
    },
    "Название вопроса 3":{
        "speakers":["Докладчик 1", "Докладчик 2"],
        "decissions":{
            "Принятое решение 1":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            },
            "Принятое решение 2":{
                "context":"краткий контекст обсуждения по достигнутому решению.",
                "time":"мм:сс"
            }
        }
    }
}
poruch={
    "1":{
        "ispolnitely":["Организация-исполнитель 1 (И. О. Фамилия руководителя)"],
        "todo":"Сделать то-то.",
        "date":"до 1 сентября 2024"
    },
    "2":{
        "ispolnitely":["Организация-исполнитель 1 (И. О. Фамилия руководителя)",
                       "Организация-исполнитель 2 (И. О. Фамилия руководителя)",
                       "Организация-исполнитель 3 (И. О. Фамилия руководителя)"],
        "todo":"Сделать то-то.",
        "date":"до 1 сентября 2024"
    }
}


paragraph.add_run("ПРОТОКОЛ").bold=True
paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
par1=doc.add_paragraph()
run = par1.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par1.add_run("совещания по теме "+theme).bold=True
par1.alignment=WD_ALIGN_PARAGRAPH.CENTER
par2=doc.add_paragraph()
run = par2.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par2.add_run("Москва").bold=True
par2.alignment=WD_ALIGN_PARAGRAPH.CENTER
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.add_run("от "+cur+ " №                 ").bold=True
par3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.add_run("Присутствовали:")
par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
for i in range(count_of_members):
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.add_run("                        ")
    par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
k=1
u=1
for i in questions:
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.add_run(voc[k]+".     "+i)
    par3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.add_run("("+", ".join(questions[i]["speakers"])+")")
    par3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    u=1
    for j in questions[i]["decissions"]:
        par3=doc.add_paragraph()
        run = par3.add_run()
        font = run.font
        font.name = 'Arial'  # Указываем имя шрифта
        font.size = Pt(12)
        par3.add_run(str(u)+".     "+j)
        par3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        par3=doc.add_paragraph()
        run = par3.add_run()
        context="Контекст обсуждения: "+str(questions[i]["decissions"][j]["context"])
        font = run.font
        font.name = 'Arial'  # Указываем имя шрифта
        font.size = Pt(12)
        par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
        left_part, right_part = context.split(":", 1)
        left_part=left_part+":"
        par3.add_run(left_part).bold=True
        par3.add_run(right_part)
        par3=doc.add_paragraph()
        run = par3.add_run()
        context=context="Время: "+str(questions[i]["decissions"][j]["time"])
        font = run.font
        font.name = 'Arial'  # Указываем имя шрифта
        font.size = Pt(12)
        par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
        left_part, right_part = context.split(":", 1)
        left_part=left_part+":"
        par3.add_run(left_part).bold=True
        par3.add_run(right_part)
        u+=1
    k+=1
doc.add_page_break()
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
par3.add_run("№                                      ")
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
par3.add_run(str(date.day)+" "+str(months[date.month-1])+" "+str(date.year)+" г.")
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.alignment=WD_ALIGN_PARAGRAPH.CENTER
par3.add_run("ПЕРЕЧЕНЬ ПОРУЧЕНИЙ").bold=True
par3=doc.add_paragraph()
run = par3.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)
par3.alignment=WD_ALIGN_PARAGRAPH.CENTER
par3.add_run("по итогам совещания по теме "+theme).bold=True
for i in poruch:
    text=str(i)+".     "
    text+=', '.join(poruch[i]["ispolnitely"])
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par3.add_run(text)
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par3.add_run(poruch[i]["todo"])
    context="Срок - "+poruch[i]["date"]
    left_part, right_part = context.split("-", 1)
    par3=doc.add_paragraph()
    run = par3.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par3.add_run(left_part+" - ").bold=True
    par3.add_run(right_part)
    



doc.save('ofic_prot.docx')
if isPassword:
    word = win32com.client.Dispatch('Word.Application')
    text=os.getcwd()
    text=text+"\\"
    doc = word.Documents.Open(text+'ofic_prot.docx')
    doc.Password = password
    doc.Save()
    word.Quit()