import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import os

import subtitle_parser

doc=Document()

paragraph = doc.add_paragraph()

password="12345"
isPassword=True

run = paragraph.add_run()
font = run.font
font.name = 'Arial'  # Указываем имя шрифта
font.size = Pt(12)

with open('output.srt', 'r', encoding="utf-8") as input_file:
    parser = subtitle_parser.SrtParser(input_file)
    parser.parse()

parser.print_warnings()
prev=""
for subtitle in parser.subtitles:
    left_part, right_part = subtitle.text.split(":", 1)
    if prev!=left_part:
        paragraph.add_run("\n").bold =True
        paragraph.add_run(left_part).bold =True
        paragraph.add_run(right_part)
        prev=left_part
    else:
        paragraph.add_run(right_part)
    

doc.save('simple_demo.docx')
if isPassword:
    word = win32com.client.Dispatch('Word.Application')
    text=os.getcwd()
    text=text+"\\"
    doc = word.Documents.Open(text+'simple_demo.docx')
    doc.Password = password
    doc.Save()
    word.Quit()