import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import msoffcrypto

from local_lib import get_chat_metadata, seconds_to_time

from pydub import (
    AudioSegment
)

def word(a:str):
    if a=="speakers":
        return "Докладчики:"
    elif a=="date":
        return "Срок:"
    elif a=="otv":
        return "Ответственный:"
    elif a=="obraz":
        return "Образ результата:"
    elif a=="context":
        return "Контекст обсуждения:"
    elif a=="time":
        return "Время:"
    else:
        return ""
    
goal="декларация цели встречи"

months = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
           'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']

voc={1:"I",2:"II",3:"III",4:"IV",5:"V",
     6:"VI",7:"VII",8:"VIII",9:"IX",10:"X",
     11:"XI",12:"XII",13:"XIII",14:"XIV",15:"XV",
     16:"XVI",17:"XVII",18:"XVIII",19:"XIX",20:"XX"}

def create_docx(data):

    metadata = get_chat_metadata(data['chat_id'])

    date=datetime.datetime.now()
    cur=str(date.day)+"."+str(date.month)+"."+str(date.year)
    sound : AudioSegment = AudioSegment.from_file(f'./temp/{metadata['chat_id']}/{metadata['audio']['filename']}')

    llm_reply = data['transcribed_text']

    resume = {}
    for i in range(len(llm_reply)):
        for j in range(len(llm_reply[i])):
            question = llm_reply[i][j]['Вопрос обсуждения']
            resume[f'{voc[j+1]}.{question}'] = {
                'b': {
                    'speakers': llm_reply[i][j]['Участники обсуждения'],
                    'decision': llm_reply[i][j]['Принятое решение'],
                    'context': llm_reply[i][j]['Контекст обсуждения'],
                    'time': llm_reply[i][j]['Время'],
                }
            }


    dur = seconds_to_time(int(sound.duration_seconds))
    members = ", ".join(metadata['members'])

    doc=Document()

    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(16)
    par1.add_run("ТЕМА ВСТРЕЧИ").bold=True
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(14)
    par1.add_run("ПРОТОКОЛ ВСТРЕЧИ")
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par1.add_run("ДАТА:        ").bold=True
    par1.add_run(cur)
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par1.add_run("ВРЕМЯ:        ").bold=True
    par1.add_run(str(date.hour)+":"+str(date.minute))
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par1.add_run("ДЛИТЕЛЬНОСТЬ:        ").bold=True
    par1.add_run(dur)
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(12)
    par1.add_run("УЧАСТНИКИ:        ").bold=True
    par1.add_run(members)
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(14)
    par1.add_run("ПОВЕСТКА ДНЯ")
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(14)
    par1.add_run("Цель встречи:").bold=True
    par1.add_run(goal)
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for i in resume:
        par1=doc.add_paragraph()
        run = par1.add_run()
        font = run.font
        font.name = 'Arial'  # Указываем имя шрифта
        font.size = Pt(12)
        par1.add_run(i)
        par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    par1=doc.add_paragraph()
    run = par1.add_run()
    font = run.font
    font.name = 'Arial'  # Указываем имя шрифта
    font.size = Pt(14)
    par1.add_run("РЕЗЮМЕ ОБСУЖДЕНИЙ")
    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for i in resume:
        par1=doc.add_paragraph()
        run = par1.add_run()
        font = run.font
        font.name = 'Arial'  # Указываем имя шрифта
        font.size = Pt(12)
        par1.add_run(i)
        par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
        for j in resume[i]:
            par1=doc.add_paragraph()
            run = par1.add_run()
            font = run.font
            font.name = 'Arial'  # Указываем имя шрифта
            font.size = Pt(12)
            par1.add_run(j+")")
            par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
            for k in resume[i][j]:
                if k=="speakers":
                    par1=doc.add_paragraph()
                    run = par1.add_run()
                    font = run.font
                    font.name = 'Arial'  # Указываем имя шрифта
                    font.size = Pt(12)
                    par1.add_run(word(k)+', '.join(resume[i][j][k]))
                    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT
                else:
                    par1=doc.add_paragraph()
                    run = par1.add_run()
                    font = run.font
                    font.name = 'Arial'  # Указываем имя шрифта
                    font.size = Pt(12)
                    par1.add_run(word(k)+(resume[i][j][k] or " "))
                    par1.alignment=WD_ALIGN_PARAGRAPH.LEFT

    filename = ''.join(data["file_name"].split(".")[:-1])
    docx_filepath = f'./temp/{data["chat_id"]}/{filename}.docx'

    doc.save(docx_filepath)
    if 'password' in metadata:
        with open(docx_filepath, 'r+b') as docx_file:
            doc = msoffcrypto.OfficeFile(docx_file)
            doc.encrypt(metadata['password'], docx_file)
            
    return docx_filepath